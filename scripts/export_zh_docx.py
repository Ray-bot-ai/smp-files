#!/usr/bin/env python3
"""全库中文合并 Word 导出：每 500 件一卷（8 卷）。
文本规则与 PDF 版一致：zh 优先；zh 空且 native_zh/is_chinese → 用 en（原文）；
其余无译文页 → 〔本页无译文〕。
用法：uv run --with python-docx scripts/export_zh_docx.py [--limit N]
"""
import glob
import json
import os
import re
import sys

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
import demo_run as D

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
OUT = os.path.join(ROOT, "export_docx")
os.makedirs(OUT, exist_ok=True)
VOL_SIZE = 500


def set_cjk(run, font="宋体", size=None, bold=None, color=None):
    run.font.name = font
    r = run._element.rPr.rFonts
    r.set(qn("w:eastAsia"), font)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor(*color)


def page_text(q):
    zh = (q.get("zh") or "").strip()
    if zh:
        return zh, False
    en = (q.get("en") or "").strip()
    if not en:
        return "", True
    if q.get("zh_status") == "native_zh" or D.is_chinese(en):
        return en, False
    return "", True


files = []
for f in glob.glob(f"{ROOT}/data/smpa-files-*.json"):
    m = re.search(r"smpa-files-(\d+)\.json$", f)
    if m:
        files.append((int(m.group(1)), f))
files.sort()

if "--limit" in sys.argv:
    files = files[:int(sys.argv[sys.argv.index("--limit") + 1])]
    print(f"[冒烟] 仅处理前 {len(files)} 件")

total_pages_written = 0
vol_no = 0
for i in range(0, len(files), VOL_SIZE):
    vol_no += 1
    chunk = files[i:i + VOL_SIZE]
    lo = re.search(r"smpa-files-(\d+)", os.path.basename(chunk[0][1])).group(1)
    hi = re.search(r"smpa-files-(\d+)", os.path.basename(chunk[-1][1])).group(1)

    doc = Document()
    # 卷标题
    h = doc.add_heading("", level=0)
    set_cjk(h.add_run(f"卷 {vol_no}｜档案 smpa-files-{lo} 至 {hi}"), "黑体", 20, True)
    p = doc.add_paragraph()
    set_cjk(p.add_run("说明：正文为中文译本；原文即中文的页直接录原文；无译文的页以〔本页无译文〕标出。"),
            "宋体", 10, color=(0x66, 0x66, 0x66))

    for idx, path in chunk:
        iid = os.path.basename(path)[:-5]
        try:
            d = json.load(open(path, encoding="utf-8"))
        except Exception:
            continue
        title = d.get("title") or ""
        n_pages = len(d.get("page_data", []))
        h = doc.add_heading("", level=2)
        set_cjk(h.add_run(f"{iid}｜{title}｜共 {n_pages} 页"), "黑体", 13, True)

        for q in sorted(d.get("page_data", []), key=lambda x: x.get("n", 0)):
            text, missing = page_text(q)
            ph = doc.add_paragraph()
            set_cjk(ph.add_run(f"{iid}｜第 {q.get('n', 0) + 1} 页"), "宋体", 9,
                    color=(0x80, 0x80, 0x80))
            if missing:
                pm = doc.add_paragraph()
                set_cjk(pm.add_run("〔本页无译文〕"), "宋体", 10.5,
                        color=(0x99, 0x99, 0x99))
            else:
                for para in text.split("\n"):
                    para = para.strip()
                    if para:
                        pp = doc.add_paragraph()
                        pp.paragraph_format.space_after = Pt(2)
                        set_cjk(pp.add_run(para), "宋体", 10.5)
        total_pages_written += n_pages

    out = os.path.join(OUT, f"卷{vol_no:02d}_smpa-files-{lo}-{hi}.docx")
    if os.path.exists(out) and "--limit" not in sys.argv:
        print(f"卷{vol_no:02d}: 已存在，跳过（断点续跑）")
        total_pages_written += sum(
            len(json.load(open(p, encoding="utf-8")).get("page_data", []))
            for _, p in chunk)
        continue
    doc.save(out)
    print(f"卷{vol_no:02d}: {len(chunk)} 件 → {out} "
          f"({os.path.getsize(out) / 1e6:.1f} MB)")

print(f"完成：{vol_no} 卷，{len(files)} 件 / {total_pages_written:,} 页")
