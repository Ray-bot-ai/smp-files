#!/usr/bin/env python3
"""把 data/*.json 里已有的中文译文导出成一个 Word 文档。

用法：
    .venv/bin/python export_docx.py                 # 全量
    .venv/bin/python export_docx.py --limit 10      # 先试 10 件，看排版
    .venv/bin/python export_docx.py -o 别的名字.docx

设计上的两条原则，改脚本时别丢：

1) 缺译文的页**不静默跳过**。这个项目已经在「静默缺口」上栽过两次
   （转录漏了不补、翻译漏了整件跳过）。导出环节同样是缺口高发区：
   一页没译文，如果直接不输出，读者会以为原件就到这儿了。所以每一页
   都要出现，没译文就写明是哪一类没有：
     · 原件本来就是中文 → 出转录原文，并注明「非译文」
     · 满页 □（模型自己说读不出）→ 注明比例
     · 译文被退化检测判废 → 注明判废理由
     · 连英文转录都没有 → 注明原件空白或转录失败
   人工认定的编造页（unusable_pages.json）同样逐页标注。

2) 不可用页只加说明、不删内容。跟 build_site.py 一个态度：
   史料产出是不可替换的，宁可留着并说清楚，也不要替读者做删除决定。

体量提示：全量约 360 万字、Word 里两三千页。python-docx 是一次性把整棵
XML 树建在内存里再写盘。注意 make_appender 那段注释里的 O(n²) 坑。
"""

import argparse
import glob
import json
import os
import re
import time

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from docx.text.paragraph import Paragraph

HERE = os.path.dirname(os.path.abspath(__file__))
DATA = os.path.join(HERE, "data")

BOX_RATIO = 0.30        # 与 build_site.py / demo_run.py 一致，改一处要改三处
EAST_ASIA = "宋体"
LATIN = "Times New Roman"
GRAY = RGBColor(0x80, 0x80, 0x80)


# ── 判据：跟 demo_run.py 保持同一套，免得两边对同一页有不同结论 ──────────

def is_chinese(text):
    """这一页本身就是中文原件吗？

    **必须与 demo_run.is_chinese 保持一致**。这里另存一份，是为了让本脚本
    不依赖 common.py 的 API key——导出 Word 不该要密钥。改一边就要改另一边。

    含假名即非中文：日文公文汉字密集，只看汉字占比会被判成中文，
    于是日文页在这里被标成「（原件即中文，未经翻译）」——把日文说成中文。
    demo_run 那边 2026-08-20 已修，这一份当时漏了（同一判据存两处的典型代价）。"""
    t = re.sub(r"\s", "", text or "")
    if len(t) < 30:
        return False
    if re.search(r"[぀-ヿㇰ-ㇿ]", text or ""):
        return False
    return sum(1 for c in t if "一" <= c <= "鿿") / len(t) > 0.5


def box_ratio(en):
    en = (en or "").strip()
    return en.count("□") / len(en) if en else 0.0


def load_manual_flags():
    p = os.path.join(HERE, "unusable_pages.json")
    if not os.path.exists(p):
        return {}
    return {(r["ia_id"], r["n"]): r["reason"]
            for r in json.load(open(p, encoding="utf-8"))}


# ── Word 细节 ──────────────────────────────────────────────────────────────

def make_appender(doc):
    """返回一个「加一个空段落」的函数，绕开 python-docx 的 add_paragraph。

    别改回 doc.add_paragraph()：它是 **O(n²)**。python-docx 每加一段都调
    insert_element_before，后者用 find() 从头扫一遍 body 去定位末尾的
    <w:sectPr>。段落数一多就爆炸——实测 5k/10k/20k 段耗时 0.34/1.14/4.40 秒
    （翻一倍涨 4 倍），本文档 25 万段落跑了 11 分钟还没完。
    这里改成抓住 sectPr 节点、直接 addprevious，是 O(1)，同样三档
    0.17/0.36/0.71 秒（翻一倍涨 2 倍），全量一分钟出头。
    """
    sect = doc.element.body.find(qn("w:sectPr"))

    def new_p(style=None):
        p = OxmlElement("w:p")
        sect.addprevious(p)
        par = Paragraph(p, doc._body)
        if style:
            par.style = style
        return par

    return new_p


def set_font(run, size=None, bold=None, color=None, latin=LATIN):
    """中文字体必须单独写 w:eastAsia，只设 run.font.name 是不生效的。"""
    run.font.name = latin
    run._element.rPr.rFonts.set(qn("w:eastAsia"), EAST_ASIA)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color


NEWP = None     # 由 main() 装上 make_appender(doc) 的产物


def heading(text, level=1, size=15, page_break=False):
    h = NEWP(f"Heading {level}")
    if page_break:
        h.paragraph_format.page_break_before = True
    set_font(h.add_run(text), size=size, bold=True)
    return h


def page_break():
    NEWP().add_run().add_break(WD_BREAK.PAGE)


def para(doc, text="", size=10.5, bold=False, color=None, align=None,
         style=None, space_after=4, indent_first=False):
    p = NEWP(style)
    if text:
        set_font(p.add_run(text), size=size, bold=bold, color=color)
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.line_spacing = 1.3
    if align is not None:
        p.alignment = align
    if indent_first:
        pf.first_line_indent = Pt(size * 2)
    return p


def add_page_number_footer(section):
    """页脚居中页码。python-docx 没有现成 API，塞一个 PAGE 域进去。"""
    from docx.oxml import OxmlElement
    p = section.footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    for tag, attr, text in (("w:fldChar", ("w:fldCharType", "begin"), None),
                            ("w:instrText", ("xml:space", "preserve"), " PAGE "),
                            ("w:fldChar", ("w:fldCharType", "end"), None)):
        el = OxmlElement(tag)
        el.set(qn(attr[0]), attr[1])
        if text:
            el.text = text
        run._r.append(el)
    set_font(run, size=9, color=GRAY)


def add_hyperlink(paragraph, url, text, size=9):
    """回原件影像的链接。引用前必回原件，所以这条链接是要点不是装饰。"""
    from docx.oxml import OxmlElement
    rid = paragraph.part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True)
    link = OxmlElement("w:hyperlink")
    link.set(qn("r:id"), rid)
    run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    for tag, val in (("w:color", "0563C1"), ("w:u", "single")):
        el = OxmlElement(tag)
        el.set(qn("w:val"), "single" if tag == "w:u" else val)
        rPr.append(el)
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), LATIN)
    fonts.set(qn("w:hAnsi"), LATIN)
    fonts.set(qn("w:eastAsia"), EAST_ASIA)
    rPr.append(fonts)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(int(size * 2)))
    rPr.append(sz)
    run.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    run.append(t)
    link.append(run)
    paragraph._p.append(link)


def setup_styles(doc):
    st = doc.styles["Normal"]
    st.font.name = LATIN
    st.font.size = Pt(10.5)
    st.element.rPr.rFonts.set(qn("w:eastAsia"), EAST_ASIA)
    for name, size in (("Heading 1", 15), ("Heading 2", 12)):
        s = doc.styles[name]
        s.font.name = LATIN
        s.font.size = Pt(size)
        s.font.bold = True
        s.font.color.rgb = RGBColor(0, 0, 0)
        s.element.rPr.rFonts.set(qn("w:eastAsia"), EAST_ASIA)


# ── 每页的处置 ─────────────────────────────────────────────────────────────

def page_body(q, iid, manual):
    """返回 (段落列表, 提示语或 None, 计数分类)。

    段落列表为空表示这页没有可读正文，只剩提示语。
    """
    zh = (q.get("zh") or "").strip()
    en = (q.get("en") or "").strip()
    flag = manual.get((iid, q["n"]))

    if flag:                       # 人工认定的编造页：内容照留，前面挂警告
        body = zh or (en if is_chinese(en) else "")
        return (split(body), "【此页不可用】" + flag, "人工标记不可用")

    if zh:
        if q.get("zh_note"):
            return split(zh), "（" + q["zh_note"] + "；以下为转录原文，非译文）", "原件即中文"
        return split(zh), None, "有译文"

    if is_chinese(en):
        return (split(en), "（原件即中文，未经翻译；以下为转录原文，非译文）", "原件即中文")

    r = box_ratio(en)
    if r >= BOX_RATIO:
        return ([], f"【本页无译文】约 {round(r * 100)}% 的字模型无法辨认（转录中标为 □），"
                    f"未送翻译。请看原件影像。", "满页□")

    if q.get("zh_status"):
        return ([], f"【本页无译文】译文被判废：{q['zh_status']}。请看原件影像。", "译文判废")

    if not en:
        return [], "【本页无内容】英文转录也是空的（原件空白页，或转录失败）。", "无转录"

    # 试译过但没留下 zh_status。实见的两页都是**英文转录本身**已经退化
    # （一页是 6/6/6/… 重复循环，一页 19% 是 □，恰在 30% 阈值以下），
    # 译文因此每次都被退化检测挡掉。说清这一层，别写「原因不明」。
    if q.get("zh_tries"):
        return ([], f"【本页无译文】翻译尝试 {q['zh_tries']} 次，产出均未通过退化检测——"
                    f"通常是英文转录本身已经退化（重复循环，或大段 □）。请看原件影像。",
                "转录退化·译不出")

    return [], "【本页无译文】原因不明，翻译环节未产出。请看原件影像。", "无译文·原因不明"


def split(text):
    return [ln.strip() for ln in (text or "").split("\n") if ln.strip()]


# ── 主流程 ─────────────────────────────────────────────────────────────────

def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("-o", "--out", default="SMP档案中文译文汇编.docx")
    ap.add_argument("--limit", type=int, default=0, help="只导前 N 件（试排版用）")
    args = ap.parse_args()

    t0 = time.time()
    manual = load_manual_flags()
    files = sorted(glob.glob(os.path.join(DATA, "*.json")),
                   key=lambda p: int(re.search(r"(\d+)\.json$", p).group(1)))
    if args.limit:
        files = files[:args.limit]

    docs = []
    for f in files:
        d = json.load(open(f, encoding="utf-8"))
        if any((q.get("zh") or "").strip() or (q.get("en") or "").strip()
               for q in d.get("page_data", [])):
            docs.append(d)

    doc = Document()
    setup_styles(doc)
    global NEWP
    NEWP = make_appender(doc)
    sec = doc.sections[0]
    add_page_number_footer(sec)

    stats = {"件": len(docs), "页": 0, "字": 0}
    kinds = {}
    for d in docs:
        for q in d.get("page_data", []):
            stats["页"] += 1
            stats["字"] += len((q.get("zh") or "").strip())

    # ── 封面 ──
    for _ in range(6):
        para(doc, "")
    para(doc, "上海公共租界工部局警务处档案", size=24, bold=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=8)
    para(doc, "中文译文汇编", size=20, bold=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)
    para(doc, "Shanghai Municipal Police Files (NARA M1750, RG263)",
         size=11, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=40)
    para(doc, f"共 {stats['件']} 件卷宗 · {stats['页']:,} 页 · 译文约 {stats['字']:,} 字",
         size=11, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
    para(doc, f"导出日期 {time.strftime('%Y 年 %m 月 %d 日')}",
         size=10, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER)

    page_break()

    # ── 凡例 ──
    heading("凡例", 1)
    for t in [
        "一、本文档是机器产物，不是校订本。每一页的处理链条是："
        "原件扫描影像 → 视觉大模型转录成英文文本 → 大模型译成中文。"
        "两道工序都会出错，且错误常常是通顺的、看不出来的。",

        "二、因此：任何一句话在写进论文之前，都必须回原件影像核对。"
        "每件卷宗标题下方附有该件在 Internet Archive 的地址，"
        "点进去可逐页看原件。本文档的用途是检索和通读，不是引用依据。",

        "三、没有译文的页，本文档一律照实标出原因，不会悄悄跳过。"
        "已知有四类：原件本来就是中文（出转录原文，并注明非译文）；"
        "满页 □（模型自己承认读不出）；译文被退化检测判废；"
        "连英文转录都没有（原件空白页或转录失败）。",

        "四、另有一类最危险的错误：模型编造。转录出的文字通顺、像模像样，"
        "但与原件内容毫无关系。这类错误任何自动判据都抓不到，只能靠人眼发现，"
        "发现一页记一页，在正文中以「此页不可用」标出。"
        "没有被标出，不等于没有问题。",

        "五、页码用原件的页序（第 0 页起，与 Internet Archive 的影像页序一致），"
        "不是 Word 的页码。引用时请报卷宗号与原件页序。",
    ]:
        para(doc, t, size=10.5, space_after=10, indent_first=True)

    page_break()

    # ── 卷宗清单 ──
    heading(f"卷宗清单（共 {len(docs)} 件）", 1)
    para(doc, "Word 左侧「视图 → 导航窗格」可按卷宗标题直接跳转。",
         size=9.5, color=GRAY, space_after=12)
    for i, d in enumerate(docs, 1):
        n_zh = sum(1 for q in d.get("page_data", []) if (q.get("zh") or "").strip())
        line = (f"{i:>3}. [{d['ia_id'].replace('smpa-files-', '')}] "
                f"{d.get('series') or '未标系列'} / NARA {d.get('nara_file_no')} · "
                f"{d.get('title', '')}（{n_zh} 页）")
        para(doc, line, size=9.5, space_after=2)

    page_break()

    # ── 正文 ──
    for i, d in enumerate(docs, 1):
        iid = d["ia_id"]
        short = iid.replace("smpa-files-", "")
        heading(f"{i}. [{short}] {d.get('title', '')}", 1, page_break=True)

        meta = f"{d.get('series') or '未标系列'} · NARA 卷宗号 {d.get('nara_file_no')} · 共 {d.get('pages')} 页"
        para(doc, meta, size=9.5, color=GRAY, space_after=2)
        p = NEWP()
        p.paragraph_format.space_after = Pt(14)
        set_font(p.add_run("原件影像："), size=9, color=GRAY)
        add_hyperlink(p, d.get("ia_url", ""), d.get("ia_url", ""))

        for q in sorted(d.get("page_data", []), key=lambda x: x["n"]):
            lines, note, kind = page_body(q, iid, manual)
            kinds[kind] = kinds.get(kind, 0) + 1

            hp = NEWP()
            hp.paragraph_format.space_before = Pt(10)
            hp.paragraph_format.space_after = Pt(2)
            set_font(hp.add_run(f"── 第 {q['n']} 页 ──"), size=9, bold=True, color=GRAY)

            if note:
                para(doc, note, size=9.5, color=RGBColor(0xA0, 0x30, 0x30), space_after=4)
            # 正文不缩进首行：译文是逐行转录式的短行（表头、编号、签名），
            # 缩进会让整页看着参差。凡例那种成段的说明文字才缩进。
            for ln in lines:
                para(doc, ln, size=10.5, space_after=4)

    out = os.path.join(HERE, args.out)
    doc.save(out)
    mb = os.path.getsize(out) / 2**20

    print(f"→ {out}")
    print(f"   {stats['件']} 件 / {stats['页']:,} 页 / 译文 {stats['字']:,} 字 / "
          f"{mb:.1f} MB / 用时 {time.time() - t0:.0f}s")
    print("   逐页处置分类：")
    for k, v in sorted(kinds.items(), key=lambda x: -x[1]):
        print(f"     {v:>6,}  {k}")


if __name__ == "__main__":
    main()
